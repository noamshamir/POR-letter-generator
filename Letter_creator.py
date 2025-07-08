import pandas as pd
from docx import Document
import datetime
from copy import deepcopy

# The following file paths are relative to the repo root
doc_template_path = 'in/Letter_Template.docx'
excel_data_path = 'in/Letter_Excel.xlsx'
output_dir = 'out'

data = pd.read_excel(excel_data_path)
data_dicts = data.to_dict(orient='records')

for index, row in enumerate(data_dicts):
    last_name = row['LAST NAME']
    first_name = row['CLIENT NAME']
    room = str(row['RM'])
    full_name = f'{first_name} {last_name}'
    caps_full_name = str.upper(full_name)
    try:
        date = row['DATE'].date()
        date_day_of_month = date.strftime('%d')
        if date_day_of_month.startswith('0'):
            date_day_of_month = date_day_of_month[1:]
        date = date.strftime(f'%b {date_day_of_month}, %Y')
    except Exception as e:
        print(f"Row {index+1}: Cannot parse date for {first_name} {last_name} ({row['DATE']}). Skipping")
        continue
    
    if not isinstance(last_name, str) or not last_name:
        print(f"Row {index+1}: Missing or invalid last name for {first_name}. Skipping")
        continue
    
    if not isinstance(first_name, str) or not first_name:
        print(f"Row {index+1}: Missing or invalid first name for {last_name}. Skipping")
        continue
    
    try:
        int_room = int(room)
    except:
        print(f"Warning Row {index+1}: Room is not a single integer. Room number: {room}. Client: {full_name}")
        room = room[ :3]

    if not isinstance(room, str) or not room:
        print(f"Row {index+1}: Missing or invalid Room Number for {last_name}. Skipping")
        continue

    new_doc = deepcopy(Document(doc_template_path))
    
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            if '<caps_full_name>' in run.text:
                run.text = run.text.replace('<caps_full_name>', caps_full_name)
            if '<full_name>' in run.text:
                run.text = run.text.replace('<full_name>', full_name)
            if '<rm>' in run.text:
                run.text = run.text.replace('<rm>', room)
            if '<date>' in run.text:
                run.text = run.text.replace('<date>', date)

    for table in new_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '<caps_full_name>' in run.text:
                            run.text = run.text.replace('<caps_full_name>', full_name)
                        if '<full_name>' in run.text:
                            run.text = run.text.replace('<full_name>', full_name)
                        if '<rm>' in run.text:
                            run.text = run.text.replace('<rm>', rm)
                        if '<date>' in run.text:
                            run.text = run.text.replace('<date>', date)

    output_filename = f'{output_file_path}{full_name}.docx'
    new_doc.save(output_filename)
